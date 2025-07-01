import base64
import io
from jinja2 import Environment, FileSystemLoader
from weasyprint import HTML
import fitz  # PyMuPDF
import copy
import re
import docx
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.shared import qn
from docx.enum.dml import MSO_THEME_COLOR_INDEX

class DocxUtils:
    @staticmethod
    def clean_na_values(data):
        """
        Recursively clean 'NA', 'N/A', empty strings, and None values from resume data.
        Returns a cleaned copy of the data.
        """
        if isinstance(data, dict):
            cleaned = {}
            for key, value in data.items():
                cleaned_value = DocxUtils.clean_na_values(value)
                # Only include the field if it has a valid value
                if cleaned_value is not None and cleaned_value != '':
                    cleaned[key] = cleaned_value
            return cleaned
        elif isinstance(data, list):
            cleaned_list = []
            for item in data:
                cleaned_item = DocxUtils.clean_na_values(item)
                # Only include the item if it's not empty after cleaning
                if cleaned_item is not None and cleaned_item != '':
                    if isinstance(cleaned_item, dict) and cleaned_item:  # Non-empty dict
                        cleaned_list.append(cleaned_item)
                    elif not isinstance(cleaned_item, dict):  # Non-dict items
                        cleaned_list.append(cleaned_item)
            return cleaned_list
        elif isinstance(data, str):
            # Clean string values
            cleaned_str = data.strip()
            # Filter out various "NA" representations
            na_values = {'na', 'n/a', 'not applicable', 'not available', 'none', 'null', '-', ''}
            if cleaned_str.lower() in na_values:
                return None
            return cleaned_str
        else:
            # Return other types as-is (numbers, booleans, etc.)
            return data

    @staticmethod
    def get_base64_image(image_path):
        with open(image_path, "rb") as img_file:
            return base64.b64encode(img_file.read()).decode("utf-8")

    @staticmethod
    def get_base64_pdf(pdf_file):
        pdf_file.seek(0)
        return base64.b64encode(pdf_file.read()).decode("utf-8")

    @staticmethod
    def clean_html_text(text):
        """Clean HTML tags and return text with bold formatting info"""
        if not text:
            return []
        
        text = str(text)
        bold_parts = []
        
        # Find all <strong> tags and their content
        strong_pattern = r'<strong>(.*?)</strong>'
        matches = list(re.finditer(strong_pattern, text, re.IGNORECASE))
        
        if matches:
            current_pos = 0
            for match in matches:
                # Add text before bold part
                if match.start() > current_pos:
                    bold_parts.append((text[current_pos:match.start()], False))
                # Add bold part
                bold_parts.append((match.group(1), True))
                current_pos = match.end()
            # Add remaining text
            if current_pos < len(text):
                bold_parts.append((text[current_pos:], False))
        else:
            bold_parts = [(text, False)]
        
        # Clean any remaining HTML tags
        cleaned_parts = []
        for part_text, is_bold in bold_parts:
            clean_text = re.sub(r'<[^>]+>', '', part_text)
            if clean_text.strip():
                cleaned_parts.append((clean_text, is_bold))
        
        return cleaned_parts

    @staticmethod
    def add_formatted_text(paragraph, text_parts, font_size=10, font_color=RGBColor(34, 34, 34)):
        """Add text with mixed formatting to a paragraph"""
        for text, is_bold in text_parts:
            if text.strip():
                run = paragraph.add_run(text)
                run.font.name = 'Montserrat'
                run.font.size = Pt(font_size)
                run.font.color.rgb = font_color
                if is_bold:
                    run.bold = True

    @staticmethod
    def add_section_title(container, title, margin_top=18):
        """Add a section title with consistent formatting matching PDF template"""
        para = container.add_paragraph()
        para.paragraph_format.space_before = Pt(margin_top)
        para.paragraph_format.space_after = Pt(4)
        run = para.add_run(title.upper())
        run.bold = True
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(242, 93, 93)  # #f25d5d
        run.font.name = 'Montserrat'
        return para

    @staticmethod
    def add_triangle_bullet_point(container, text_parts, indent=22):
        """Add bullet point with triangle (▶) matching PDF template"""
        para = container.add_paragraph()
        para.paragraph_format.left_indent = Pt(indent)
        para.paragraph_format.space_after = Pt(6)
        
        # Add red triangle
        triangle_run = para.add_run("▶ ")
        triangle_run.font.name = 'Montserrat'
        triangle_run.font.size = Pt(13)
        triangle_run.font.color.rgb = RGBColor(242, 93, 93)
        triangle_run.bold = True
        
        # Add text content
        DocxUtils.add_formatted_text(para, text_parts, font_size=12)
        return para

    @staticmethod
    def add_robust_page_border(doc):
        """Add page border with enhanced compatibility for Word web/desktop and PDF export"""
        try:
            for section in doc.sections:
                sectPr = section._sectPr
                pgBorders = OxmlElement('w:pgBorders')
                pgBorders.set(qn('w:offsetFrom'), 'page')
                pgBorders.set(qn('w:display'), 'allPages')
                
                # Add border on all sides with standard settings for better compatibility
                for border_name in ['top', 'left', 'bottom', 'right']:
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), 'single')
                    border.set(qn('w:sz'), '6')  # Slightly thicker for better visibility in PDF
                    border.set(qn('w:space'), '0')  # Closer to page edge for consistency
                    border.set(qn('w:color'), 'F25D5D')  # Resume red color
                    pgBorders.append(border)
                
                sectPr.append(pgBorders)
        except Exception as e:
            print(f"Could not add page border: {e}")

    @staticmethod
    def create_compatible_table(parent, rows, cols, width_inches=None):
        """Create a table with enhanced compatibility for Word web/desktop"""
        try:
            # Check if parent is a header/footer (which requires width parameter)
            if hasattr(parent, '_sectPr') or 'header' in str(type(parent)).lower() or 'footer' in str(type(parent)).lower():
                # Header/footer tables require width parameter
                if width_inches is None:
                    width_inches = 8.0  # Default width for headers/footers
                table = parent.add_table(rows=rows, cols=cols, width=Inches(width_inches))
            else:
                # Regular document tables
                table = parent.add_table(rows=rows, cols=cols)
        except Exception as e:
            # Fallback: try with width parameter
            try:
                if width_inches is None:
                    width_inches = 8.0
                table = parent.add_table(rows=rows, cols=cols, width=Inches(width_inches))
            except:
                # Last resort: use basic add_table without parameters
                table = parent.add_table(rows, cols)
        
        # Configure table properties
        try:
            table.autofit = False
            table.allow_autofit = False
        except:
            pass  # Some table types might not support these properties
        
        # Apply compatible table properties
        try:
            tbl = table._tbl
            tblPr = tbl.tblPr
            
            # Use fixed layout to maintain column widths during PDF export
            tblLayout = OxmlElement('w:tblLayout')
            tblLayout.set(qn('w:type'), 'fixed')
            tblPr.append(tblLayout)
        except:
            pass  # Fallback gracefully
        
        return table

    @staticmethod
    def add_column_border(cell, border_side='right', color='D3D3D3', width='6'):
        """Add a border to a specific side of a table cell with fallback compatibility"""
        try:
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = tcPr.find(qn('w:tcBorders'))
            if tcBorders is None:
                tcBorders = OxmlElement('w:tcBorders')
                tcPr.append(tcBorders)
            
            border = OxmlElement(f'w:{border_side}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), width)  # Border thickness
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), color)
            tcBorders.append(border)
        except Exception as e:
            print(f"Could not add cell border: {e}")

    @staticmethod
    def remove_all_table_borders(table):
        """Remove all table borders for seamless layout with better compatibility"""
        try:
            tbl = table._tbl
            tblPr = tbl.tblPr
            
            # Remove table borders
            tblBorders = OxmlElement('w:tblBorders')
            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'nil')
                tblBorders.append(border)
            tblPr.append(tblBorders)
            
            # Remove cell borders
            for row in table.rows:
                for cell in row.cells:
                    try:
                        tc = cell._tc
                        tcPr = tc.get_or_add_tcPr()
                        tcBorders = OxmlElement('w:tcBorders')
                        for border_name in ['top', 'left', 'bottom', 'right']:
                            border = OxmlElement(f'w:{border_name}')
                            border.set(qn('w:val'), 'nil')
                            tcBorders.append(border)
                        tcPr.append(tcBorders)
                    except:
                        continue  # Skip if border removal fails
        except Exception as e:
            print(f"Could not remove table borders: {e}")

    @staticmethod
    def set_standard_spacing(paragraph, space_before_pt=0, space_after_pt=0):
        """Set paragraph spacing using standard methods for better compatibility"""
        if space_before_pt > 0:
            paragraph.paragraph_format.space_before = Pt(space_before_pt)
        if space_after_pt > 0:
            paragraph.paragraph_format.space_after = Pt(space_after_pt)

    @staticmethod
    def apply_standard_font(run, font_name='Montserrat', font_size_pt=10, is_bold=False, color_rgb=None):
        """Apply font formatting using standard methods for better compatibility"""
        run.font.name = font_name
        run.font.size = Pt(font_size_pt)
        if is_bold:
            run.bold = True
        if color_rgb:
            run.font.color.rgb = color_rgb

    @staticmethod
    def add_cell_background_compatible(cell, color_hex='F2F2F2'):
        """Add cell background with enhanced compatibility"""
        try:
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), color_hex)
            cell._tc.get_or_add_tcPr().append(shading_elm)
            
            # Add standard cell margins for consistent spacing
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_mar = OxmlElement('w:tcMar')
            
            # Set margins using standard values
            for margin in ['top', 'left', 'bottom', 'right']:
                mar_elem = OxmlElement(f'w:{margin}')
                mar_elem.set(qn('w:w'), '144')  # 144 twentieths = 10pt
                mar_elem.set(qn('w:type'), 'dxa')
                tc_mar.append(mar_elem)
            
            tc_pr.append(tc_mar)
            return True
        except Exception as e:
            print(f"Could not add cell background: {e}")
            return False

    @staticmethod
    def create_compatible_footer_table(footer, width_inches=8.5):
        """Create a footer table with enhanced compatibility"""
        try:
            footer_table = footer.add_table(rows=1, cols=1)
            footer_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            footer_table.autofit = False
            
            # Set table width
            footer_table.columns[0].width = Inches(width_inches)
            
            # Remove borders for seamless appearance
            DocxUtils.remove_all_table_borders(footer_table)
            
            return footer_table
        except Exception as e:
            print(f"Could not create footer table: {e}")
            return None
    
    @staticmethod
    def ensure_table_column_borders(table, column_index=0, border_color='CCCCCC'):
        """Ensure table has proper column borders for clear separation"""
        try:
            if column_index < len(table.columns) - 1:  # Don't add border to last column
                for row in table.rows:
                    cell = row.cells[column_index]
                    DocxUtils.add_column_border(cell, 'right', border_color, '8')
        except Exception as e:
            print(f"Could not add column borders: {e}")
    
    @staticmethod
    def set_fixed_column_widths(table, left_width_inches, right_width_inches):
        """Set fixed column widths that will be maintained during PDF export"""
        try:
            # Set precise column widths
            table.columns[0].width = Inches(left_width_inches)
            table.columns[1].width = Inches(right_width_inches)
            
            # Force fixed table layout
            tbl = table._tbl
            tblPr = tbl.tblPr
            
            # Ensure fixed layout
            existing_layout = tblPr.find(qn('w:tblLayout'))
            if existing_layout is not None:
                tblPr.remove(existing_layout)
            
            tblLayout = OxmlElement('w:tblLayout')
            tblLayout.set(qn('w:type'), 'fixed')
            tblPr.append(tblLayout)
            
            # Set table width for consistency
            tblW = OxmlElement('w:tblW')
            tblW.set(qn('w:w'), str(int((left_width_inches + right_width_inches) * 1440)))  # Convert to twentieths
            tblW.set(qn('w:type'), 'dxa')
            tblPr.append(tblW)
            
        except Exception as e:
            print(f"Could not set fixed column widths: {e}")

    @staticmethod
    def lock_all_table_layouts(doc):
        """Lock all table layouts to fixed for consistent PDF export"""
        try:
            # Find all tables in the document and set them to fixed layout
            for table in doc.tables:
                try:
                    tbl = table._tbl
                    tblPr = tbl.tblPr
                    
                    # Remove existing layout if present
                    existing_layout = tblPr.find(qn('w:tblLayout'))
                    if existing_layout is not None:
                        tblPr.remove(existing_layout)
                    
                    # Set to fixed layout
                    tblLayout = OxmlElement('w:tblLayout')
                    tblLayout.set(qn('w:type'), 'fixed')
                    tblPr.append(tblLayout)
                except:
                    continue  # Skip if table cannot be processed
            
            # Also check header and footer tables
            for section in doc.sections:
                try:
                    # Header tables
                    if section.header:
                        for table in section.header.tables:
                            try:
                                tbl = table._tbl
                                tblPr = tbl.tblPr
                                existing_layout = tblPr.find(qn('w:tblLayout'))
                                if existing_layout is not None:
                                    tblPr.remove(existing_layout)
                                tblLayout = OxmlElement('w:tblLayout')
                                tblLayout.set(qn('w:type'), 'fixed')
                                tblPr.append(tblLayout)
                            except:
                                continue
                    
                    # Footer tables
                    if section.footer:
                        for table in section.footer.tables:
                            try:
                                tbl = table._tbl
                                tblPr = tbl.tblPr
                                existing_layout = tblPr.find(qn('w:tblLayout'))
                                if existing_layout is not None:
                                    tblPr.remove(existing_layout)
                                tblLayout = OxmlElement('w:tblLayout')
                                tblLayout.set(qn('w:type'), 'fixed')
                                tblPr.append(tblLayout)
                            except:
                                continue
                except:
                    continue
        except Exception as e:
            print(f"Could not lock table layouts: {e}")

    @staticmethod
    def optimize_for_pdf_export(doc):
        """Apply optimizations for better PDF export from Word"""
        try:
            # Ensure consistent font embedding and layout
            for section in doc.sections:
                # Set print layout optimizations
                section.different_first_page_header_footer = False
                section.start_type = WD_SECTION.NEW_PAGE
                
                # Ensure consistent page setup
                section.page_width = Inches(8.5)
                section.page_height = Inches(11)
        except Exception as e:
            print(f"Could not optimize for PDF export: {e}")

    @staticmethod
    def add_background_watermark(doc, bg_image_path="templates/bg.png"):
        """Add background watermark matching PDF template opacity"""
        try:
            section = doc.sections[0]
            header = section.header
            
            # Create a paragraph for the watermark
            watermark_para = header.add_paragraph()
            watermark_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            try:
                run = watermark_para.add_run()
                # Add picture with size matching template (35% of page)
                picture = run.add_picture(bg_image_path, width=Inches(2.8))
                
                # Access the drawing element to set transparency (17% opacity like template)
                drawing_elements = run._element.xpath('.//wp:anchor | .//wp:inline')
                if drawing_elements:
                    drawing = drawing_elements[0]
                    pic_elements = drawing.xpath('.//pic:pic')
                    if pic_elements:
                        pic = pic_elements[0]
                        blip_elements = pic.xpath('.//a:blipFill')
                        if blip_elements:
                            blipFill = blip_elements[0]
                            
                            # Add alpha modulation for transparency (17% opacity)
                            alphaModFix = OxmlElement('a:alphaModFix')
                            alphaModFix.set('amt', '17000')  # 17% opacity (17000 out of 100000)
                            
                            # Find or create the effect list
                            effectLst = blipFill.find('.//a:effectLst')
                            if effectLst is None:
                                effectLst = OxmlElement('a:effectLst')
                                blipFill.append(effectLst)
                            
                            effectLst.append(alphaModFix)
                
            except Exception as e:
                print(f"Could not add background image watermark: {e}")
            
            # Position the watermark behind text
            watermark_para.paragraph_format.space_before = Pt(0)
            watermark_para.paragraph_format.space_after = Pt(0)
            
            return True
        except Exception as e:
            print(f"Could not add watermark background: {e}")
            return False

    # Keep old method names for backward compatibility but use new implementations
    @staticmethod
    def add_grey_background(cell):
        """Add grey background to cell - compatibility wrapper"""
        return DocxUtils.add_cell_background_compatible(cell, 'F2F2F2')

    @staticmethod
    def remove_table_borders(table):
        """Remove table borders - compatibility wrapper"""
        return DocxUtils.remove_all_table_borders(table)
    
    @staticmethod
    def add_page_border(doc):
        """Add page border - compatibility wrapper"""
        return DocxUtils.add_robust_page_border(doc)
    
    @staticmethod
    def optimize_table_for_word(table):
        """Optimize table for Word - compatibility wrapper"""
        return DocxUtils.remove_all_table_borders(table)
    
    @staticmethod
    def add_word_optimized_spacing(paragraph, space_before=0, space_after=0, line_spacing=1.0):
        """Add spacing - compatibility wrapper"""
        return DocxUtils.set_standard_spacing(paragraph, space_before, space_after)
    
    @staticmethod
    def add_word_font_optimization(run, font_name='Montserrat', font_size=10, is_bold=False, color_rgb=None):
        """Font optimization - compatibility wrapper"""
        return DocxUtils.apply_standard_font(run, font_name, font_size, is_bold, color_rgb)

    @staticmethod
    def generate_docx(data, keywords=None, left_logo_path="templates/left_logo_small.png", right_logo_path="templates/right_logo_small.png"):
        """
        Generate a .docx resume matching the PDF template exactly.
        Returns a BytesIO object containing the Word file.
        """

        # Create a deep copy and clean NA values
        data_copy = copy.deepcopy(data)
        data_copy = DocxUtils.clean_na_values(data_copy)
        
        # Limit skills to prevent left column overflow (max 18 skills for DOCX)
        if data_copy.get('skills') and len(data_copy['skills']) > 18:
            data_copy['skills'] = data_copy['skills'][:18]
        
        # Limit certifications to prevent overflow (max 5 certifications)
        if data_copy.get('certifications') and len(data_copy['certifications']) > 5:
            data_copy['certifications'] = data_copy['certifications'][:5]
        
        doc = docx.Document()
        
        # Set page margins with small values for compatibility
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.2)    # Small margin for compatibility
            section.bottom_margin = Inches(0.2)
            section.left_margin = Inches(0.2)
            section.right_margin = Inches(0.2)
            
            # Set header and footer distances for better compatibility
            section.header_distance = Inches(0.15)
            section.footer_distance = Inches(0.15)
            
            # Standard page setup for better compatibility
            section.page_width = Inches(8.5)   # Standard letter width
            section.page_height = Inches(11)   # Standard letter height

        # Add robust page border for better compatibility
        DocxUtils.add_robust_page_border(doc)
        
        # Skip watermark for now as requested
        # DocxUtils.add_background_watermark(doc)

        # --- HEADER WITH LOGOS (Compatible design) ---
        header = doc.sections[0].header
        header.is_linked_to_previous = False
        
        # Clear any default header content
        for para in header.paragraphs:
            try:
                p = para._element
                p.getparent().remove(p)
            except:
                para.clear()  # Fallback method
        
        # Create header table with compatible design
        header_table = DocxUtils.create_compatible_table(header, rows=1, cols=3, width_inches=8.1)
        
        # Set column widths for balanced layout (proportional to page)
        header_table.columns[0].width = Inches(2.7)
        header_table.columns[1].width = Inches(2.7)
        header_table.columns[2].width = Inches(2.7)
        
        # Left logo - optimized for Word
        left_cell = header_table.cell(0, 0)
        left_para = left_cell.paragraphs[0]
        left_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        DocxUtils.add_word_optimized_spacing(left_para, space_before=2, space_after=0)
        left_para.paragraph_format.left_indent = Pt(12)
        try:
            left_run = left_para.add_run()
            left_run.add_picture(left_logo_path, height=Inches(0.35))
        except Exception:
            left_run = left_para.add_run("ShorthillsAI")
            DocxUtils.add_word_font_optimization(left_run, 'Montserrat', 10, True, RGBColor(242, 93, 93))
        
        # Right logo - optimized for Word
        right_cell = header_table.cell(0, 2)
        right_para = right_cell.paragraphs[0]
        right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        DocxUtils.add_word_optimized_spacing(right_para, space_before=2, space_after=0)
        right_para.paragraph_format.right_indent = Pt(12)
        try:
            right_run = right_para.add_run()
            right_run.add_picture(right_logo_path, height=Inches(0.45))
        except Exception:
            right_run = right_para.add_run("Microsoft Partner")
            right_run.font.name = 'Montserrat'
            right_run.font.size = Pt(10)
            right_run.font.color.rgb = RGBColor(102, 102, 102)
            
        # Remove spacing paragraph after header
        # spacing_para = doc.add_paragraph()
        # DocxUtils.add_word_optimized_spacing(spacing_para, space_after=1)

        # --- MAIN CONTENT TABLE (Compatible two-column layout) ---
        main_table = DocxUtils.create_compatible_table(doc, rows=1, cols=2, width_inches=8.1)
        
        # Set fixed column widths: 35% left, 65% right (based on 8.1" usable width)
        DocxUtils.set_fixed_column_widths(main_table, 2.8, 5.3)  # 35% and 65% split
        
        # Add clear vertical border between columns for better separation
        DocxUtils.ensure_table_column_borders(main_table, 0, 'CCCCCC')
        
        left_cell = main_table.cell(0, 0)
        right_cell = main_table.cell(0, 1)
        
        # Clear default paragraphs
        left_cell._tc.clear_content()
        right_cell._tc.clear_content()

        # Add smaller padding to left cell
        left_cell_padding = left_cell.add_paragraph()
        left_cell_padding.paragraph_format.left_indent = Pt(12)

        # --- LEFT COLUMN (Compatible design) ---
        # Name and Title container with grey background
        name_title_table = DocxUtils.create_compatible_table(left_cell, rows=2, cols=1, width_inches=2.8)
        name_title_table.columns[0].width = Inches(2.8)  # Match left column width (35%)
        
        # Name cell - Word optimized
        name_cell = name_title_table.cell(0, 0)
        DocxUtils.add_grey_background(name_cell)
        name_para = name_cell.paragraphs[0]
        DocxUtils.add_word_optimized_spacing(name_para, space_before=0, space_after=2)
        name_para.paragraph_format.left_indent = Pt(8)
        name_run = name_para.add_run(data_copy.get('name', ''))
        DocxUtils.add_word_font_optimization(name_run, 'Montserrat', 17, True, RGBColor(242, 93, 93))
        
        # Title cell - Word optimized
        title_cell = name_title_table.cell(1, 0)
        DocxUtils.add_grey_background(title_cell)
        title_para = title_cell.paragraphs[0]
        DocxUtils.add_word_optimized_spacing(title_para, space_after=10)
        title_para.paragraph_format.left_indent = Pt(8)
        title_parts = DocxUtils.clean_html_text(data_copy.get('title', ''))
        for text, is_bold in title_parts:
            if text.strip():
                title_run = title_para.add_run(text)
                DocxUtils.add_word_font_optimization(title_run, 'Montserrat', 13, True, RGBColor(34, 34, 34))

        # Skills Section - Word optimized
        if data_copy.get('skills'):
            skills_para = left_cell.add_paragraph()
            DocxUtils.add_word_optimized_spacing(skills_para, space_before=10, space_after=2)
            skills_para.paragraph_format.left_indent = Pt(12)
            skills_run = skills_para.add_run('SKILLS')
            DocxUtils.add_word_font_optimization(skills_run, 'Montserrat', 11, True, RGBColor(242, 93, 93))
            
            for skill in data_copy['skills']:
                skill_parts = DocxUtils.clean_html_text(skill)
                skill_para = left_cell.add_paragraph()
                DocxUtils.add_word_optimized_spacing(skill_para, space_after=3)
                skill_para.paragraph_format.left_indent = Pt(28)
                
                # Add red triangle with Word optimization
                arrow_run = skill_para.add_run("▶ ")
                DocxUtils.add_word_font_optimization(arrow_run, 'Montserrat', 11, True, RGBColor(242, 93, 93))
                
                # Add skill text with Word optimization
                for text, is_bold in skill_parts:
                    if text.strip():
                        skill_run = skill_para.add_run(text)
                        DocxUtils.add_word_font_optimization(skill_run, 'Montserrat', 10, is_bold, RGBColor(34, 34, 34))

        # Education Section - Word optimized
        if data_copy.get('education'):
            edu_para = left_cell.add_paragraph()
            DocxUtils.add_word_optimized_spacing(edu_para, space_before=10, space_after=2)
            edu_para.paragraph_format.left_indent = Pt(12)
            edu_run = edu_para.add_run('EDUCATION')
            DocxUtils.add_word_font_optimization(edu_run, 'Montserrat', 11, True, RGBColor(242, 93, 93))
            
            for edu in data_copy['education']:
                if isinstance(edu, dict):
                    para = left_cell.add_paragraph()
                    DocxUtils.add_word_optimized_spacing(para, space_after=3)
                    para.paragraph_format.left_indent = Pt(28)
                    
                    # Add red triangle
                    arrow_run = para.add_run("▶ ")
                    DocxUtils.add_word_font_optimization(arrow_run, 'Montserrat', 11, True, RGBColor(242, 93, 93))
                    
                    # Add degree
                    if edu.get('degree'):
                        degree_parts = DocxUtils.clean_html_text(edu['degree'])
                        for text, is_bold in degree_parts:
                            if text.strip():
                                degree_run = para.add_run(text)
                                DocxUtils.add_word_font_optimization(degree_run, 'Montserrat', 10, is_bold, RGBColor(34, 34, 34))
                    
                    # Add institution on new line
                    if edu.get('institution'):
                        para.add_run('\n')
                        inst_parts = DocxUtils.clean_html_text(edu['institution'])
                        for text, is_bold in inst_parts:
                            if text.strip():
                                inst_run = para.add_run(text)
                                DocxUtils.add_word_font_optimization(inst_run, 'Montserrat', 10, True, RGBColor(34, 34, 34))

        # Certifications Section - Word optimized
        if data_copy.get('certifications'):
            cert_para = left_cell.add_paragraph()
            DocxUtils.add_word_optimized_spacing(cert_para, space_before=10, space_after=2)
            cert_para.paragraph_format.left_indent = Pt(12)
            cert_run = cert_para.add_run('CERTIFICATIONS')
            DocxUtils.add_word_font_optimization(cert_run, 'Montserrat', 11, True, RGBColor(242, 93, 93))
            
            for cert in data_copy['certifications']:
                para = left_cell.add_paragraph()
                DocxUtils.add_word_optimized_spacing(para, space_after=3)
                para.paragraph_format.left_indent = Pt(28)
                
                # Add red triangle
                arrow_run = para.add_run("▶ ")
                DocxUtils.add_word_font_optimization(arrow_run, 'Montserrat', 11, True, RGBColor(242, 93, 93))
                
                if isinstance(cert, dict):
                    # Track if we have any content to add
                    has_content = False
                    
                    # Add certification title
                    if cert.get('title'):
                        title_parts = DocxUtils.clean_html_text(cert['title'])
                        for text, is_bold in title_parts:
                            if text.strip():
                                title_run = para.add_run(text)
                                DocxUtils.add_word_font_optimization(title_run, 'Montserrat', 10, is_bold, RGBColor(34, 34, 34))
                                has_content = True
                    
                    # Add issuer (only if we have title or other content)
                    if cert.get('issuer') and has_content:
                        para.add_run('\n')
                        issuer_parts = DocxUtils.clean_html_text(cert['issuer'])
                        for text, is_bold in issuer_parts:
                            if text.strip():
                                issuer_run = para.add_run(text)
                                DocxUtils.add_word_font_optimization(issuer_run, 'Montserrat', 10, True, RGBColor(34, 34, 34))
                    elif cert.get('issuer') and not has_content:
                        # If no title but have issuer, add issuer as main content
                        issuer_parts = DocxUtils.clean_html_text(cert['issuer'])
                        for text, is_bold in issuer_parts:
                            if text.strip():
                                issuer_run = para.add_run(text)
                                DocxUtils.add_word_font_optimization(issuer_run, 'Montserrat', 10, True, RGBColor(34, 34, 34))
                                has_content = True
                    
                    # Add year (only if we have other content)
                    if cert.get('year') and has_content:
                        year_run = para.add_run(f"\n{cert['year']}")
                        DocxUtils.add_word_font_optimization(year_run, 'Montserrat', 10, False, RGBColor(34, 34, 34))

        # --- RIGHT COLUMN (Word-optimized) ---
        # Add right column padding
        right_cell_padding = right_cell.add_paragraph()
        right_cell_padding.paragraph_format.left_indent = Pt(12)

        # Summary Section - Word optimized
        if data_copy.get('summary'):
            summary_title_para = right_cell.add_paragraph()
            DocxUtils.add_word_optimized_spacing(summary_title_para, space_before=0, space_after=2)
            summary_title_para.paragraph_format.left_indent = Pt(12)
            summary_title_run = summary_title_para.add_run('SUMMARY')
            DocxUtils.add_word_font_optimization(summary_title_run, 'Montserrat', 11, True, RGBColor(242, 93, 93))
            
            summary_para = right_cell.add_paragraph()
            DocxUtils.add_word_optimized_spacing(summary_para, space_after=5)
            summary_para.paragraph_format.left_indent = Pt(12)
            summary_parts = DocxUtils.clean_html_text(data_copy['summary'])
            for text, is_bold in summary_parts:
                if text.strip():
                    summary_run = summary_para.add_run(text)
                    DocxUtils.add_word_font_optimization(summary_run, 'Montserrat', 10, is_bold, RGBColor(34, 34, 34))

        # Projects Section - Word optimized
        if data_copy.get('projects'):
            # Add minimal spacing
            spacing_para = right_cell.add_paragraph()
            DocxUtils.add_word_optimized_spacing(spacing_para, space_after=3)
            
            # Section title
            section_para = right_cell.add_paragraph()
            DocxUtils.add_word_optimized_spacing(section_para, space_after=2)
            section_para.paragraph_format.left_indent = Pt(12)
            section_run = section_para.add_run('KEY RESPONSIBILITIES:')
            DocxUtils.add_word_font_optimization(section_run, 'Montserrat', 11, True, RGBColor(242, 93, 93))
            
            # Add minimal spacing
            spacing_para2 = right_cell.add_paragraph()
            DocxUtils.add_word_optimized_spacing(spacing_para2, space_after=1)
            
            # Project entries
            for idx, project in enumerate(data_copy['projects']):
                # Project title
                if project.get('title'):
                    proj_title_para = right_cell.add_paragraph()
                    DocxUtils.add_word_optimized_spacing(proj_title_para, space_before=6, space_after=2)
                    proj_title_para.paragraph_format.left_indent = Pt(12)
                    
                    title_run = proj_title_para.add_run(f"Project {idx + 1}: ")
                    DocxUtils.add_word_font_optimization(title_run, 'Montserrat', 11, True, RGBColor(242, 93, 93))
                    
                    title_parts = DocxUtils.clean_html_text(project['title'])
                    for text, is_bold in title_parts:
                        if text.strip():
                            proj_run = proj_title_para.add_run(text)
                            DocxUtils.add_word_font_optimization(proj_run, 'Montserrat', 11, True, RGBColor(242, 93, 93))
                
                # Project description bullets
                if project.get('description'):
                    desc_text = project['description']
                    if isinstance(desc_text, str):
                        # Enhanced bullet splitting to handle HTML lists, newlines, and bullet characters
                        desc_text = desc_text.replace('</li>', '</li>\n').replace('•', '\n')
                        # Split by full stops, newlines, <br>, or <li>.
                        bullets = re.split(r'(?<=[.?!])\s+|\n|<br\s*/?>|<li>', desc_text)
                        
                        for bullet_html in bullets:
                            bullet_html = bullet_html.strip()
                            if not bullet_html:
                                continue
                            
                            bullet_parts = DocxUtils.clean_html_text(bullet_html)
                            if not any(part[0].strip() for part in bullet_parts):
                                continue

                            bullet_para = right_cell.add_paragraph()
                            DocxUtils.add_word_optimized_spacing(bullet_para, space_after=3)
                            bullet_para.paragraph_format.left_indent = Pt(28)
                            
                            # Add red triangle
                            arrow_run = bullet_para.add_run("▶ ")
                            DocxUtils.add_word_font_optimization(arrow_run, 'Montserrat', 11, True, RGBColor(242, 93, 93))
                            
                            # Add bullet text
                            for text, is_bold in bullet_parts:
                                if text.strip():
                                    bullet_run = bullet_para.add_run(text)
                                    DocxUtils.add_word_font_optimization(bullet_run, 'Montserrat', 10, is_bold, RGBColor(34, 34, 34))
        
        # --- FOOTER (Compatible design) ---
        footer = doc.sections[0].footer
        footer.is_linked_to_previous = False
        
        # Clear any default footer content
        for para in footer.paragraphs:
            try:
                p = para._element
                p.getparent().remove(p)
            except:
                para.clear()  # Fallback method
            
        # Create footer table with compatible design
        footer_table = DocxUtils.create_compatible_footer_table(footer, 8.0)
        if footer_table is not None:
            footer_cell = footer_table.cell(0, 0)
            
            # Set cell background color to brand orange
            DocxUtils.add_cell_background_compatible(footer_cell, 'F25D5D')
            
            # Add footer text with proper formatting
            footer_para = footer_cell.paragraphs[0]
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            DocxUtils.set_standard_spacing(footer_para, space_before_pt=6, space_after_pt=6)
            
            footer_run = footer_para.add_run("© www.shorthills.ai")
            DocxUtils.apply_standard_font(footer_run, 'Montserrat', 10, False, RGBColor(255, 255, 255))
        
        # Apply PDF export optimizations and lock column widths
        DocxUtils.optimize_for_pdf_export(doc)
        DocxUtils.lock_all_table_layouts(doc)
        
        # Save document to BytesIO
        docx_file = io.BytesIO()
        doc.save(docx_file)
        docx_file.seek(0)
        return docx_file

    @staticmethod
    def generate_multi_page_docx(data, keywords=None):
        """
        Generate a multi-page .docx resume with proper page breaks and section distribution.
        This handles large amounts of content that don't fit on a single page.
        """

        def estimate_content_size(data):
            """Estimate how much content we have to determine if multi-page is needed"""
            size = 0
            size += len(data.get('skills', [])) * 2  # Each skill takes ~2 units
            size += len(data.get('education', [])) * 4  # Each education entry ~4 units
            size += len(data.get('certifications', [])) * 3  # Each cert ~3 units
            
            # Projects can be large
            for project in data.get('projects', []):
                project_size = 5  # Base size for title
                if project.get('description'):
                    bullets = project['description'].split('\n') if isinstance(project['description'], str) else []
                    project_size += len([b for b in bullets if b.strip()]) * 2
                size += project_size
            
            return size

        # Apply keyword bolding if keywords provided and clean NA values
        data_copy = copy.deepcopy(data) if keywords else data
        data_copy = DocxUtils.clean_na_values(data_copy)
        
        # Limit skills to prevent left column overflow (max 18 skills for DOCX)
        if data_copy.get('skills') and len(data_copy['skills']) > 18:
            data_copy['skills'] = data_copy['skills'][:18]
        
        # Limit certifications to prevent overflow (max 5 certifications)
        if data_copy.get('certifications') and len(data_copy['certifications']) > 5:
            data_copy['certifications'] = data_copy['certifications'][:5]
        
        # Estimate if we need multiple pages
        content_size = estimate_content_size(data_copy)
        needs_multiple_pages = content_size > 65  # Threshold for single page
        
        if not needs_multiple_pages:
            # Use single page layout
            return DocxUtils.generate_docx(data_copy, keywords)
        
        # For multi-page, create a modified single-page layout that flows naturally
        # Word will handle page breaks automatically with proper styling
        doc = docx.Document()
        
        # Set page margins with small values for compatibility
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.2)
            section.bottom_margin = Inches(0.2)
            section.left_margin = Inches(0.2)
            section.right_margin = Inches(0.2)

            # Set header and footer distances for compatibility
            section.header_distance = Inches(0.15)
            section.footer_distance = Inches(0.15)

        # Add robust page border for better compatibility
        DocxUtils.add_robust_page_border(doc)
        # Skip watermark as requested
        # DocxUtils.add_background_watermark(doc)

        # Header with logos (will appear on all pages)
        header = doc.sections[0].header
        
        # Clear any default header content
        for para in header.paragraphs:
            para.clear()
            
        header_table = DocxUtils.create_compatible_table(header, rows=1, cols=3, width_inches=8.1)
        header_table.columns[0].width = Inches(2.7)
        header_table.columns[1].width = Inches(2.7)
        header_table.columns[2].width = Inches(2.7)
        
        # Left logo
        left_cell = header_table.cell(0, 0)
        left_para = left_cell.paragraphs[0]
        left_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        left_para.paragraph_format.space_before = Pt(8)
        left_para.paragraph_format.space_after = Pt(4)
        left_para.paragraph_format.left_indent = Pt(8)
        try:
            left_para.add_run().add_picture("templates/left_logo_small.png", height=Inches(0.5))
        except Exception:
            left_run = left_para.add_run("ShorthillsAI")
            left_run.font.name = 'Montserrat'
            left_run.font.size = Pt(12)
            left_run.font.color.rgb = RGBColor(242, 93, 93)
            left_run.bold = True
        
        # Right logo
        right_cell = header_table.cell(0, 2)
        right_para = right_cell.paragraphs[0]
        right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        right_para.paragraph_format.space_before = Pt(8)
        right_para.paragraph_format.space_after = Pt(4)
        right_para.paragraph_format.right_indent = Pt(8)
        try:
            right_para.add_run().add_picture("templates/right_logo_small.png", height=Inches(0.6))
        except Exception:
            right_run = right_para.add_run("Microsoft Partner")
            right_run.font.name = 'Montserrat'
            right_run.font.size = Pt(10)
            right_run.font.color.rgb = RGBColor(102, 102, 102)

        DocxUtils.remove_all_table_borders(header_table)

        # For multi-page, use a simplified linear layout instead of complex 2-column
        # This ensures better page breaks and readability
        
        # Add small padding for content
        content_padding = doc.add_paragraph()
        content_padding.paragraph_format.left_indent = Pt(12)
        content_padding.paragraph_format.space_after = Pt(6)

        # Name and title (centered)
        name_para = doc.add_paragraph()
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_para.paragraph_format.space_before = Pt(20)
        name_para.paragraph_format.space_after = Pt(6)
        name_run = name_para.add_run(data_copy.get('name', ''))
        name_run.bold = True
        name_run.font.size = Pt(24)
        name_run.font.color.rgb = RGBColor(242, 93, 93)
        name_run.font.name = 'Montserrat'
        
        if data_copy.get('title'):
            title_para = doc.add_paragraph()
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_para.paragraph_format.space_after = Pt(20)
            title_parts = DocxUtils.clean_html_text(data_copy['title'])
            for text, is_bold in title_parts:
                if text.strip():
                    title_run = title_para.add_run(text)
                    title_run.font.size = Pt(16)
                    title_run.font.name = 'Montserrat'
                    title_run.font.color.rgb = RGBColor(34, 34, 34)
                    title_run.bold = True

        # Summary
        if data_copy.get('summary'):
            summary_title_para = doc.add_paragraph()
            summary_title_para.paragraph_format.left_indent = Pt(12)
            summary_title_para.paragraph_format.space_before = Pt(0)
            summary_title_para.paragraph_format.space_after = Pt(4)
            summary_title_run = summary_title_para.add_run('SUMMARY')
            summary_title_run.bold = True
            summary_title_run.font.size = Pt(13)
            summary_title_run.font.color.rgb = RGBColor(242, 93, 93)
            summary_title_run.font.name = 'Montserrat'
            
            summary_para = doc.add_paragraph()
            summary_para.paragraph_format.left_indent = Pt(12)
            summary_para.paragraph_format.space_after = Pt(18)
            summary_parts = DocxUtils.clean_html_text(data_copy['summary'])
            DocxUtils.add_formatted_text(summary_para, summary_parts, font_size=12)

        # Projects
        if data_copy.get('projects'):
            projects_title_para = doc.add_paragraph()
            projects_title_para.paragraph_format.left_indent = Pt(12)
            projects_title_para.paragraph_format.space_after = Pt(4)
            projects_title_run = projects_title_para.add_run('KEY RESPONSIBILITIES')
            projects_title_run.bold = True
            projects_title_run.font.size = Pt(13)
            projects_title_run.font.color.rgb = RGBColor(242, 93, 93)
            projects_title_run.font.name = 'Montserrat'
            
            for idx, project in enumerate(data_copy['projects']):
                if project.get('title'):
                    proj_title_para = doc.add_paragraph()
                    proj_title_para.paragraph_format.left_indent = Pt(12)
                    proj_title_para.paragraph_format.space_before = Pt(12)
                    proj_title_para.paragraph_format.space_after = Pt(4)
                    
                    title_run = proj_title_para.add_run(f"Project {idx + 1}: ")
                    title_run.bold = True
                    title_run.font.size = Pt(13)
                    title_run.font.color.rgb = RGBColor(242, 93, 93)
                    title_run.font.name = 'Montserrat'
                    
                    title_parts = DocxUtils.clean_html_text(project['title'])
                    for text, is_bold in title_parts:
                        if text.strip():
                            proj_run = proj_title_para.add_run(text)
                            proj_run.bold = True
                            proj_run.font.size = Pt(13)
                            proj_run.font.color.rgb = RGBColor(242, 93, 93)
                            proj_run.font.name = 'Montserrat'

                if project.get('description'):
                    desc_text = project['description']
                    if isinstance(desc_text, str):
                        # Enhanced bullet splitting to handle HTML lists, newlines, and bullet characters
                        desc_text = desc_text.replace('</li>', '</li>\n').replace('•', '\n')
                        # Split by full stops, newlines, <br>, or <li>.
                        bullets = re.split(r'(?<=[.?!])\s+|\n|<br\s*/?>|<li>', desc_text)
                        
                        for bullet_html in bullets:
                            bullet_html = bullet_html.strip()
                            if not bullet_html:
                                continue
                            
                            bullet_parts = DocxUtils.clean_html_text(bullet_html)

                            if not any(part[0].strip() for part in bullet_parts):
                                continue

                            bullet_para = doc.add_paragraph()
                            bullet_para.paragraph_format.left_indent = Pt(34)  # 12 + 22
                            bullet_para.paragraph_format.space_after = Pt(6)
                            
                            # Add red triangle
                            arrow_run = bullet_para.add_run("▶ ")
                            arrow_run.font.name = 'Montserrat'
                            arrow_run.font.size = Pt(13)
                            arrow_run.font.color.rgb = RGBColor(242, 93, 93)
                            arrow_run.bold = True
                            
                            # Add bullet text
                            DocxUtils.add_formatted_text(bullet_para, bullet_parts, font_size=12)

        # Skills
        if data_copy.get('skills'):
            skills_title_para = doc.add_paragraph()
            skills_title_para.paragraph_format.left_indent = Pt(12)
            skills_title_para.paragraph_format.space_before = Pt(18)
            skills_title_para.paragraph_format.space_after = Pt(4)
            skills_title_run = skills_title_para.add_run('SKILLS')
            skills_title_run.bold = True
            skills_title_run.font.size = Pt(13)
            skills_title_run.font.color.rgb = RGBColor(242, 93, 93)
            skills_title_run.font.name = 'Montserrat'
            
            for skill in data_copy['skills']:
                skill_para = doc.add_paragraph()
                skill_para.paragraph_format.left_indent = Pt(34)  # 12 + 22
                skill_para.paragraph_format.space_after = Pt(6)
                
                # Add red triangle
                arrow_run = skill_para.add_run("▶ ")
                arrow_run.font.name = 'Montserrat'
                arrow_run.font.size = Pt(13)
                arrow_run.font.color.rgb = RGBColor(242, 93, 93)
                arrow_run.bold = True
                
                # Add skill text
                skill_parts = DocxUtils.clean_html_text(skill)
                DocxUtils.add_formatted_text(skill_para, skill_parts, font_size=12)

        # Education
        if data_copy.get('education'):
            edu_title_para = doc.add_paragraph()
            edu_title_para.paragraph_format.left_indent = Pt(12)
            edu_title_para.paragraph_format.space_before = Pt(18)
            edu_title_para.paragraph_format.space_after = Pt(4)
            edu_title_run = edu_title_para.add_run('EDUCATION')
            edu_title_run.bold = True
            edu_title_run.font.size = Pt(13)
            edu_title_run.font.color.rgb = RGBColor(242, 93, 93)
            edu_title_run.font.name = 'Montserrat'
            
            for edu in data_copy['education']:
                if isinstance(edu, dict):
                    para = doc.add_paragraph()
                    para.paragraph_format.left_indent = Pt(34)  # 12 + 22
                    para.paragraph_format.space_after = Pt(6)
                
                    arrow_run = para.add_run("▶ ")
                    arrow_run.font.name = 'Montserrat'
                    arrow_run.font.size = Pt(13)
                    arrow_run.font.color.rgb = RGBColor(242, 93, 93)
                    arrow_run.bold = True
                    
                    if edu.get('degree'):
                        degree_parts = DocxUtils.clean_html_text(edu['degree'])
                        DocxUtils.add_formatted_text(para, degree_parts, font_size=12)
                    
                    if edu.get('institution'):
                        para.add_run('\n')
                        inst_parts = DocxUtils.clean_html_text(edu['institution'])
                        for text, is_bold in inst_parts:
                            if text.strip():
                                inst_run = para.add_run(text)
                                inst_run.font.name = 'Montserrat'
                                inst_run.font.size = Pt(12)
                                inst_run.font.color.rgb = RGBColor(34, 34, 34)
                                inst_run.bold = True
                else:
                    edu_para = doc.add_paragraph()
                    edu_para.paragraph_format.left_indent = Pt(34)
                    edu_para.paragraph_format.space_after = Pt(6)
                    
                    arrow_run = edu_para.add_run("▶ ")
                    arrow_run.font.name = 'Montserrat'
                    arrow_run.font.size = Pt(13)
                    arrow_run.font.color.rgb = RGBColor(242, 93, 93)
                    arrow_run.bold = True
                    
                    edu_parts = DocxUtils.clean_html_text(str(edu))
                    DocxUtils.add_formatted_text(edu_para, edu_parts, font_size=12)

        # Certifications
        if data_copy.get('certifications'):
            cert_title_para = doc.add_paragraph()
            cert_title_para.paragraph_format.left_indent = Pt(12)
            cert_title_para.paragraph_format.space_before = Pt(18)
            cert_title_para.paragraph_format.space_after = Pt(4)
            cert_title_run = cert_title_para.add_run('CERTIFICATIONS')
            cert_title_run.bold = True
            cert_title_run.font.size = Pt(13)
            cert_title_run.font.color.rgb = RGBColor(242, 93, 93)
            cert_title_run.font.name = 'Montserrat'
            
            for cert in data_copy['certifications']:
                para = doc.add_paragraph()
                para.paragraph_format.left_indent = Pt(34)  # 12 + 22
                para.paragraph_format.space_after = Pt(6)
                
                arrow_run = para.add_run("▶ ")
                arrow_run.font.name = 'Montserrat'
                arrow_run.font.size = Pt(13)
                arrow_run.font.color.rgb = RGBColor(242, 93, 93)
                arrow_run.bold = True
                
                if isinstance(cert, dict):
                    # Track if we have any content to add
                    has_content = False
                    
                    if cert.get('title'):
                        title_parts = DocxUtils.clean_html_text(cert['title'])
                        DocxUtils.add_formatted_text(para, title_parts, font_size=12)
                        has_content = True
                    
                    # Add issuer (only if we have title or other content)
                    if cert.get('issuer') and has_content:
                        para.add_run('\n')
                        issuer_parts = DocxUtils.clean_html_text(cert['issuer'])
                        for text, is_bold in issuer_parts:
                            if text.strip():
                                issuer_run = para.add_run(text)
                                issuer_run.font.name = 'Montserrat'
                                issuer_run.font.size = Pt(12)
                                issuer_run.font.color.rgb = RGBColor(34, 34, 34)
                                issuer_run.bold = True
                    elif cert.get('issuer') and not has_content:
                        # If no title but have issuer, add issuer as main content
                        issuer_parts = DocxUtils.clean_html_text(cert['issuer'])
                        for text, is_bold in issuer_parts:
                            if text.strip():
                                issuer_run = para.add_run(text)
                                issuer_run.font.name = 'Montserrat'
                                issuer_run.font.size = Pt(12)
                                issuer_run.font.color.rgb = RGBColor(34, 34, 34)
                                issuer_run.bold = True
                                has_content = True
                    
                    # Add year (only if we have other content)
                    if cert.get('year') and has_content:
                        year_run = para.add_run(f"\n{cert['year']}")
                        year_run.font.name = 'Montserrat'
                        year_run.font.size = Pt(12)
                        year_run.font.color.rgb = RGBColor(34, 34, 34)
                else:
                    cert_parts = DocxUtils.clean_html_text(str(cert))
                    DocxUtils.add_formatted_text(para, cert_parts, font_size=12)

        # Footer with compatible design
        footer = doc.sections[0].footer
        
        # Clear any default footer content
        for para in footer.paragraphs:
            para.clear()
            
        # Create footer table for consistent background
        footer_table = DocxUtils.create_compatible_footer_table(footer, 8.0)
        if footer_table is not None:
            footer_cell = footer_table.cell(0, 0)
            DocxUtils.add_cell_background_compatible(footer_cell, 'F25D5D')
            
            footer_para = footer_cell.paragraphs[0]
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            DocxUtils.set_standard_spacing(footer_para, space_before_pt=6, space_after_pt=6)
            
            footer_run = footer_para.add_run("© www.shorthills.ai")
            DocxUtils.apply_standard_font(footer_run, 'Montserrat', 10, False, RGBColor(255, 255, 255))

        # Save document
        docx_file = io.BytesIO()
        doc.save(docx_file)
        docx_file.seek(0)
        return docx_file